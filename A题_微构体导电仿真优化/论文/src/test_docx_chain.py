from __future__ import annotations

import json
from pathlib import Path

import pytest
from PIL import Image
from docx import Document
from docx.oxml.ns import qn

from audit_docx import Auditor
from build_docx import (
    BuildState,
    QA_IDENTITY,
    TEMPLATE_SHA256,
    add_code_block,
    add_figure_grid,
    configure_styles,
    expand_content_slots,
    load_content_map,
    load_metadata,
    materialize_content_slot,
    parse_markdown,
    sha256_file,
    validate_identity,
)
from check_docx_readiness import check_readiness
from prepare_submission_sources import prepare_submission_sources


SRC_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = SRC_DIR.parents[1]
AI_HEADER_LINES = [
    "# AI 工具：OpenAI Codex；模型/版本：GPT-5 系列；开发机构：OpenAI。",
    "# 版本发布日期：2025-08-07（GPT-5 系列公开快照日期）；本程序由参赛队逐行复核并对结果负责。",
]


def write_allowlist(
    project_root: Path,
    files: list[dict[str, object]],
    *,
    status: str = "frozen",
) -> Path:
    path = project_root / "source-allowlist.json"
    path.write_text(
        json.dumps(
            {
                "schema_version": "1.0",
                "status": status,
                "ai_header_lines": AI_HEADER_LINES,
                "files": files,
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    return path


def prepare_simple_source_project(
    project_root: Path,
    source_text: str = "from __future__ import annotations\n\nprint('ok')\n",
) -> tuple[Path, Path, dict[str, object]]:
    source = project_root / "main.py"
    source.write_text(source_text, encoding="utf-8")
    allowlist = write_allowlist(
        project_root,
        [{"path": "main.py", "title": "附录 C.1 主程序"}],
    )
    manifest = prepare_submission_sources(
        allowlist,
        project_root / "提交源码",
        project_root,
    )
    return source, project_root / "提交源码" / "main.py", manifest


def materialize_source_appendix(project_root: Path):
    return materialize_content_slot(
        "SOURCE_CODE_APPENDIX",
        {
            "type": "source_code_appendix",
            "manifest": "提交源码/source-manifest.json",
        },
        project_root,
        project_root,
    )


def test_markdown_parser_supports_required_blocks() -> None:
    source = """# 标题

## 一、正文

普通段含行内公式 $x^2$。

1. 编号项

| 列一 | 列二 |
|---|---:|
| $x$ | 1 |

$$
y=x^2
$$

![图注](figure.png)

<!-- FIGURE_SLOT -->

```python
print('ok')
```
"""
    title, blocks = parse_markdown(source)
    assert title == "标题"
    kinds = [block.kind for block in blocks]
    assert kinds == ["heading", "paragraph", "list", "table", "math", "image", "placeholder", "code"]


def test_internal_qa_identity_and_output_guard() -> None:
    allowed = PROJECT_ROOT / "tmp" / "docx_test" / "paper-internal-qa.docx"
    validate_identity(QA_IDENTITY, QA_IDENTITY, True, allowed, PROJECT_ROOT)
    with pytest.raises(ValueError):
        validate_identity(QA_IDENTITY, QA_IDENTITY, True, PROJECT_ROOT / "论文" / "paper-internal-qa.docx", PROJECT_ROOT)
    with pytest.raises(ValueError):
        validate_identity(QA_IDENTITY, QA_IDENTITY, True, PROJECT_ROOT / "tmp" / "ACM1234567.docx", PROJECT_ROOT)
    with pytest.raises(ValueError):
        validate_identity(QA_IDENTITY, QA_IDENTITY, False, allowed, PROJECT_ROOT)


def test_retained_template_hash() -> None:
    template = PROJECT_ROOT / "00_赛题与附件" / "2026第七届华数杯竞赛论文模板.docx"
    assert sha256_file(template) == TEMPLATE_SHA256


def test_metadata_file_loads_abstract_and_keywords(tmp_path: Path) -> None:
    metadata = tmp_path / "paper.yml"
    metadata.write_text("abstract: 已冻结摘要\nkeywords: [随机几何, 连通概率]\n", encoding="utf-8")
    assert load_metadata(metadata) == {
        "abstract": "已冻结摘要",
        "keywords": ["随机几何", "连通概率"],
    }


def test_formal_content_slots_resolve_from_fragments_and_source_allowlist(tmp_path: Path) -> None:
    fragments = tmp_path / "fragments"
    fragments.mkdir()
    (fragments / "ai.md").write_text("#### A.1 工具信息\n\n人工复核。\n", encoding="utf-8")
    (fragments / "commands.md").write_text(
        "#### B.1 复现命令\n\n```powershell\npython main.py\n```\n", encoding="utf-8"
    )
    source, submission, _ = prepare_simple_source_project(tmp_path)
    content_map_path = tmp_path / "content_map.json"
    content_map_path.write_text(
        json.dumps(
            {
                "AI_USAGE_DETAILS": {"type": "markdown", "path": "fragments/ai.md"},
                "REPRODUCTION_COMMANDS": {"type": "markdown", "path": "fragments/commands.md"},
                "SOURCE_CODE_APPENDIX": {
                    "type": "source_code_appendix",
                    "manifest": "提交源码/source-manifest.json",
                },
            },
            ensure_ascii=False,
        ),
        encoding="utf-8",
    )
    markdown = tmp_path / "paper.md"
    markdown.write_text(
        "# 测试论文\n\n## 附录\n\n<!-- AI_USAGE_DETAILS -->\n\n"
        "<!-- REPRODUCTION_COMMANDS -->\n\n<!-- SOURCE_CODE_APPENDIX -->\n",
        encoding="utf-8",
    )
    metadata = tmp_path / "paper.yml"
    metadata.write_text("abstract: 摘要\nkeywords: [关键词]\n", encoding="utf-8")
    figure_map = tmp_path / "figure_map.json"
    figure_map.write_text("{}\n", encoding="utf-8")

    report = check_readiness(
        markdown,
        metadata,
        figure_map,
        content_map_path,
        project_root=tmp_path,
    )
    assert report["ready"] is True
    assert report["unresolved_slots"] == []
    assert report["resolved_content_slots"] == [
        "AI_USAGE_DETAILS",
        "REPRODUCTION_COMMANDS",
        "SOURCE_CODE_APPENDIX",
    ]
    assert report["included_source_files"] == [
        {
            "slot": "SOURCE_CODE_APPENDIX",
            "title": "附录 C.1 主程序",
            "source_path": "main.py",
            "source_sha256": sha256_file(source),
            "path": "提交源码/main.py",
            "language": "python",
            "sha256": sha256_file(submission),
            "bytes": submission.stat().st_size,
        }
    ]

    _, blocks = parse_markdown(markdown.read_text(encoding="utf-8"))
    content_map, content_base = load_content_map(content_map_path)
    state = BuildState()
    expanded = expand_content_slots(blocks, content_map, content_base, tmp_path, state)
    assert all(block.kind != "placeholder" for block in expanded)
    assert [block.text for block in expanded if block.kind == "code"][-1] == submission.read_text(encoding="utf-8")
    assert state.resolved_content_slots == report["resolved_content_slots"]


def test_frozen_q1_source_hash_is_unchanged() -> None:
    source = PROJECT_ROOT / "问题" / "问题1" / "src" / "solve.py"
    assert sha256_file(source) == "6BDFE87D490EF47EDA3065BD81AE23FDD6B3DEAB484A228E8024D5E3E6CDDFA4"


def test_prepare_submission_copy_preserves_source_and_adds_exact_two_line_header(tmp_path: Path) -> None:
    source, submission, manifest = prepare_simple_source_project(tmp_path)
    original_hash = manifest["files"][0]["source_sha256"]
    assert sha256_file(source) == original_hash
    lines = submission.read_text(encoding="utf-8").splitlines()
    assert lines[:2] == AI_HEADER_LINES
    assert lines[2:] == source.read_text(encoding="utf-8").splitlines()
    assert sum(line.startswith("# AI 工具：") or line.startswith("# 版本发布日期：") for line in lines) == 2


def test_q1_submission_copy_applies_exactly_seven_controlled_replacements(tmp_path: Path) -> None:
    source = tmp_path / "问题" / "问题1" / "src" / "solve.py"
    source.parent.mkdir(parents=True)
    old_values = [f"旧报告字符串{index}" for index in range(1, 8)]
    new_values = [f"正式报告字符串{index}" for index in range(1, 8)]
    source.write_text("\n".join(f"REPORT_{index} = {value!r}" for index, value in enumerate(old_values)) + "\n", encoding="utf-8")
    original_hash = sha256_file(source)
    allowlist = write_allowlist(
        tmp_path,
        [
            {
                "path": "问题/问题1/src/solve.py",
                "title": "附录 C.1 问题一程序",
                "replacements": [
                    {"old": old, "new": new} for old, new in zip(old_values, new_values, strict=True)
                ],
                "forbidden_terms": old_values,
            }
        ],
    )
    manifest = prepare_submission_sources(allowlist, tmp_path / "提交源码", tmp_path)
    submission = tmp_path / "提交源码" / "问题" / "问题1" / "src" / "solve.py"
    submission_text = submission.read_text(encoding="utf-8")
    assert sha256_file(source) == original_hash
    assert submission_text.splitlines()[:2] == AI_HEADER_LINES
    assert all(old not in submission_text for old in old_values)
    assert all(new in submission_text for new in new_values)
    replacement_audit = manifest["files"][0]["replacement_audit"]
    assert len(replacement_audit) == 7
    assert [item["count"] for item in replacement_audit] == [1] * 7


def test_prepare_submission_sources_rejects_pending_q4(tmp_path: Path) -> None:
    (tmp_path / "main.py").write_text("print('ok')\n", encoding="utf-8")
    allowlist = write_allowlist(
        tmp_path,
        [{"path": "main.py", "title": "主程序"}],
        status="pending_q4",
    )
    with pytest.raises(ValueError, match="status must be frozen"):
        prepare_submission_sources(allowlist, tmp_path / "提交源码", tmp_path)


def test_submission_manifest_rejects_duplicate_file(tmp_path: Path) -> None:
    _, _, _ = prepare_simple_source_project(tmp_path)
    manifest_path = tmp_path / "提交源码" / "source-manifest.json"
    manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
    manifest["files"].append(dict(manifest["files"][0]))
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False), encoding="utf-8")
    with pytest.raises(ValueError, match="Duplicate source appendix path"):
        materialize_source_appendix(tmp_path)


def test_prepare_submission_sources_rejects_private_path(tmp_path: Path) -> None:
    private_text = r"print(r'C:\Users\example\secret')" + "\n"
    (tmp_path / "main.py").write_text(private_text, encoding="utf-8")
    allowlist = write_allowlist(tmp_path, [{"path": "main.py", "title": "主程序"}])
    with pytest.raises(ValueError, match="user-specific absolute path"):
        prepare_submission_sources(allowlist, tmp_path / "提交源码", tmp_path)


@pytest.mark.parametrize(
    ("drift_target", "message"),
    [
        ("source", "Original source hash changed"),
        ("submission", "Submission source hash mismatch"),
    ],
)
def test_submission_manifest_rejects_source_or_copy_hash_drift(
    tmp_path: Path,
    drift_target: str,
    message: str,
) -> None:
    source, submission, _ = prepare_simple_source_project(tmp_path)
    target = source if drift_target == "source" else submission
    target.write_text(target.read_text(encoding="utf-8") + "# drift\n", encoding="utf-8")
    with pytest.raises(ValueError, match=message):
        materialize_source_appendix(tmp_path)


def test_figure_grid_supports_two_column_two_plus_one_layout(tmp_path: Path) -> None:
    image_paths = []
    for index in range(3):
        path = tmp_path / f"figure-{index}.png"
        Image.new("RGB", (400, 300), (240 - index * 20, 245, 250)).save(path)
        image_paths.append(path)
    items = [
        {"path": path.name, "label": f"({index + 1})", "alt": f"图 {index + 1}"}
        for index, path in enumerate(image_paths)
    ]
    doc = Document()
    configure_styles(doc)
    resolved = add_figure_grid(doc, items, tmp_path, "总图注", columns=2)

    assert resolved == image_paths
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 2
    assert len(table.columns) == 2
    assert table.cell(1, 0)._tc is table.cell(1, 1)._tc
    grid_span = table.cell(1, 0)._tc.tcPr.find(qn("w:gridSpan"))
    assert grid_span is not None and grid_span.get(qn("w:val")) == "2"
    assert len(list(table._tbl.iter(qn("w:tbl")))) == 1
    assert not list(table._tbl.iter(qn("w:trHeight")))
    assert len(list(table._tbl.iter(qn("w:drawing")))) == 3
    extents = list(table._tbl.iter(qn("wp:extent")))
    assert len(extents) == 3
    assert extents[0].get("cx") == extents[2].get("cx")
    for paragraph in table.cell(1, 0)._tc.iter(qn("w:p")):
        ppr = paragraph.find(qn("w:pPr"))
        assert ppr is not None and ppr.find(qn("w:keepNext")) is not None

    body_children = list(doc._element.body)
    table_index = body_children.index(table._tbl)
    caption = body_children[table_index + 1]
    assert "".join(node.text or "" for node in caption.iter(qn("w:t"))) == "总图注"


def test_figure_grid_rejects_invalid_column_count(tmp_path: Path) -> None:
    path = tmp_path / "figure.png"
    Image.new("RGB", (40, 30), "white").save(path)
    with pytest.raises(ValueError, match="columns"):
        add_figure_grid(Document(), [{"path": path.name}], tmp_path, "", columns=0)


def test_table_geometry_audit_accepts_grid_span_in_two_plus_one_layout(tmp_path: Path) -> None:
    image_paths = []
    for index in range(3):
        path = tmp_path / f"audit-figure-{index}.png"
        Image.new("RGB", (160, 120), "white").save(path)
        image_paths.append(path)
    doc = Document()
    doc.add_table(rows=1, cols=1)  # The formal cover table is excluded from body-table checks.
    configure_styles(doc)
    add_figure_grid(
        doc,
        [{"path": path.name, "label": str(index)} for index, path in enumerate(image_paths)],
        tmp_path,
        "图注",
        columns=2,
    )
    output = tmp_path / "grid-span.docx"
    doc.save(output)

    auditor = Auditor(output, internal_qa=True)
    assert auditor.load() is True
    auditor.audit_numbering_and_tables()
    geometry = next(check for check in auditor.checks if check.id == "structure.body_table_geometry")
    assert geometry.ok is True


def test_code_block_uses_one_compact_non_grid_paragraph() -> None:
    doc = Document()
    configure_styles(doc)
    before = len(doc.paragraphs)
    add_code_block(doc, "line_1\nline_2\nline_3\n")
    assert len(doc.paragraphs) == before + 1
    paragraph = doc.paragraphs[-1]
    assert paragraph.text == "line_1\nline_2\nline_3"
    snap_to_grid = paragraph._p.get_or_add_pPr().find(qn("w:snapToGrid"))
    assert snap_to_grid is not None and snap_to_grid.get(qn("w:val")) == "0"
    style = doc.styles["PaperCode"]
    assert style.font.size.pt == pytest.approx(7.0)
    assert style.paragraph_format.line_spacing.pt == pytest.approx(8.0)
    assert style.paragraph_format.widow_control is False


def test_readiness_rejects_q4_preview_asset(tmp_path: Path) -> None:
    markdown = tmp_path / "paper.md"
    markdown.write_text("# 测试论文\n\n<!-- FIGURE_Q4_3D -->\n", encoding="utf-8")
    metadata = tmp_path / "paper.yml"
    metadata.write_text("abstract: 摘要\nkeywords: [关键词]\n", encoding="utf-8")
    preview = tmp_path / "q4_preview.png"
    preview.write_bytes(b"preview")
    figure_map = tmp_path / "figure_map.json"
    figure_map.write_text(
        json.dumps({"FIGURE_Q4_3D": {"path": preview.name}}), encoding="utf-8"
    )
    content_map = tmp_path / "content_map.json"
    content_map.write_text("{}\n", encoding="utf-8")

    report = check_readiness(
        markdown,
        metadata,
        figure_map,
        content_map,
        project_root=tmp_path,
    )
    assert report["ready"] is False
    assert report["preview_assets"] == [
        {"slot": "FIGURE_Q4_3D", "path": str(preview.resolve())}
    ]
    assert "preview_assets_forbidden" in report["issues"]


def test_internal_qa_docx_passes_but_formal_mode_rejects() -> None:
    docx = PROJECT_ROOT / "tmp" / "docx_qa" / "paper-internal-qa.docx"
    if not docx.exists():
        pytest.skip("Run the documented INTERNAL-QA build before this integration check")
    qa_report = Auditor(docx, internal_qa=True).run()
    assert qa_report["passed"] is True
    assert qa_report["blocker_failures"] == []
    formal_report = Auditor(docx, internal_qa=False).run()
    assert formal_report["passed"] is False
    assert set(formal_report["blocker_failures"]) == {"identity.slots", "content.no_unresolved_slots"}


def test_build_manifest_never_marks_internal_qa_as_formal() -> None:
    manifest = PROJECT_ROOT / "tmp" / "docx_qa" / "paper-internal-qa.build.json"
    if not manifest.exists():
        pytest.skip("Run the documented INTERNAL-QA build before this integration check")
    data = json.loads(manifest.read_text(encoding="utf-8"))
    assert data["mode"] == "internal_qa"
    assert data["formal_submission_allowed"] is False
    assert data["group"] == QA_IDENTITY
    assert data["competition_id"] == QA_IDENTITY
