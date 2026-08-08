#!/usr/bin/env python3
"""Build the competition paper from the retained official DOCX template.

The official template is copied and edited in place. The first-page table is
preserved as an OOXML object; all legacy body placeholders are replaced with
structured content parsed from the Markdown draft. LaTeX math is converted to
native Word OMML through Pandoc.
"""

from __future__ import annotations

import argparse
import copy
import hashlib
import json
import os
import re
import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Iterable, Sequence

import yaml
from lxml import etree
from PIL import Image
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, Twips


TEMPLATE_SHA256 = "4EDF51DBCDB27F842209A54E9BECAD5E1C7308CBEC3A6D7F0EDF0DBE860FC07C"
CM_PATTERN = re.compile(r"^CM[0-9]{7}$")
ALLOWED_GROUPS = {"研究生", "本科生", "专科生"}
QA_IDENTITY = "INTERNAL-QA"
USABLE_WIDTH_DXA = 9070

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
CP_NS = "http://schemas.openxmlformats.org/package/2006/metadata/core-properties"
DC_NS = "http://purl.org/dc/elements/1.1/"
EP_NS = "http://schemas.openxmlformats.org/officeDocument/2006/extended-properties"

NS = {"w": W_NS, "m": M_NS, "r": R_NS, "rel": REL_NS, "ct": CT_NS}
INLINE_MATH_RE = re.compile(r"(?<!\\)\$(?!\$)(.+?)(?<!\\)\$", re.DOTALL)
PLACEHOLDER_RE = re.compile(r"^\s*<!--\s*([A-Z][A-Z0-9_]+)\s*-->\s*$")
HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*$")
LIST_RE = re.compile(r"^(\s*)(?:(\d+)[.)]|([-+*]))\s+(.+?)\s*$")
IMAGE_RE = re.compile(r"^!\[([^\]]*)\]\(([^)]+)\)\s*$")
TABLE_SEP_CELL_RE = re.compile(r"^:?-{3,}:?$")
TOKEN_RE = re.compile(r"^[A-Z][A-Z0-9_]+$")
SOURCE_LANGUAGE_BY_SUFFIX = {
    ".c": "c",
    ".cpp": "cpp",
    ".h": "c",
    ".hpp": "cpp",
    ".java": "java",
    ".jl": "julia",
    ".js": "javascript",
    ".m": "matlab",
    ".ps1": "powershell",
    ".py": "python",
    ".r": "r",
    ".sh": "bash",
    ".sql": "sql",
    ".ts": "typescript",
}
MAX_SOURCE_FILE_BYTES = 2 * 1024 * 1024
PRIVATE_HOME_RE = re.compile(
    r"(?i)(?:[A-Z]:[\\/]+(?:Users|Documents and Settings)[\\/]+[^\\/\s\"']+|/(?:home|Users)/[^/\s\"']+)"
)


@dataclass
class Block:
    kind: str
    text: str = ""
    level: int = 0
    ordered: bool = False
    rows: list[list[str]] = field(default_factory=list)
    alignments: list[str] = field(default_factory=list)
    language: str = ""
    path: str = ""


@dataclass
class BuildState:
    unresolved: list[str] = field(default_factory=list)
    inserted_figures: list[str] = field(default_factory=list)
    resolved_content_slots: list[str] = field(default_factory=list)
    content_inputs: list[dict[str, Any]] = field(default_factory=list)
    included_sources: list[dict[str, Any]] = field(default_factory=list)
    content_errors: list[dict[str, str]] = field(default_factory=list)
    formula_count: int = 0
    numbered_paragraphs: int = 0
    bullet_paragraphs: int = 0


@dataclass
class ContentSlotResult:
    blocks: list[Block] = field(default_factory=list)
    inputs: list[dict[str, Any]] = field(default_factory=list)
    sources: list[dict[str, Any]] = field(default_factory=list)


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def smart_join(lines: Sequence[str]) -> str:
    result = ""
    for raw in lines:
        piece = raw.strip()
        if not piece:
            continue
        if result and result[-1:].isascii() and result[-1:].isalnum() and piece[:1].isascii() and piece[:1].isalnum():
            result += " "
        result += piece
    return result


def parse_front_matter(text: str) -> tuple[dict[str, Any], str]:
    if not text.startswith("---\n"):
        return {}, text
    end = text.find("\n---\n", 4)
    if end < 0:
        raise ValueError("Markdown front matter starts with --- but has no closing delimiter")
    data = yaml.safe_load(text[4:end]) or {}
    if not isinstance(data, dict):
        raise ValueError("Markdown front matter must be a mapping")
    return data, text[end + 5 :]


def split_table_row(line: str) -> list[str]:
    stripped = line.strip()
    if stripped.startswith("|"):
        stripped = stripped[1:]
    if stripped.endswith("|"):
        stripped = stripped[:-1]
    return [cell.strip().replace("\\|", "|") for cell in re.split(r"(?<!\\)\|", stripped)]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(TABLE_SEP_CELL_RE.fullmatch(cell.replace(" ", "")) for cell in cells)


def table_alignments(line: str) -> list[str]:
    result: list[str] = []
    for cell in split_table_row(line):
        value = cell.replace(" ", "")
        if value.startswith(":") and value.endswith(":"):
            result.append("center")
        elif value.endswith(":"):
            result.append("right")
        else:
            result.append("left")
    return result


def starts_special(lines: Sequence[str], index: int) -> bool:
    line = lines[index]
    if not line.strip():
        return True
    if HEADING_RE.match(line) or PLACEHOLDER_RE.match(line) or LIST_RE.match(line):
        return True
    if IMAGE_RE.match(line.strip()) or line.lstrip().startswith("```") or line.strip().startswith("$$"):
        return True
    return index + 1 < len(lines) and "|" in line and is_table_separator(lines[index + 1])


def parse_markdown(text: str) -> tuple[str, list[Block]]:
    lines = text.replace("\r\n", "\n").replace("\r", "\n").split("\n")
    blocks: list[Block] = []
    title = ""
    index = 0
    in_comment = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()
        if in_comment:
            if "-->" in stripped:
                in_comment = False
            index += 1
            continue
        if stripped.startswith("<!--") and not stripped.endswith("-->"):
            in_comment = True
            index += 1
            continue
        if not stripped:
            index += 1
            continue

        placeholder = PLACEHOLDER_RE.match(line)
        if placeholder:
            blocks.append(Block(kind="placeholder", text=placeholder.group(1)))
            index += 1
            continue
        if stripped.startswith("<!--") and stripped.endswith("-->"):
            index += 1
            continue

        heading = HEADING_RE.match(line)
        if heading:
            level = len(heading.group(1))
            value = heading.group(2).strip()
            if level == 1 and not title:
                title = value
            else:
                blocks.append(Block(kind="heading", text=value, level=level))
            index += 1
            continue

        if stripped.startswith("```"):
            language = stripped[3:].strip()
            index += 1
            code: list[str] = []
            while index < len(lines) and not lines[index].lstrip().startswith("```"):
                code.append(lines[index])
                index += 1
            if index >= len(lines):
                raise ValueError("Unclosed fenced code block")
            index += 1
            blocks.append(Block(kind="code", text="\n".join(code), language=language))
            continue

        if stripped.startswith("$$"):
            if stripped != "$$" and stripped.endswith("$$") and len(stripped) > 4:
                formula = stripped[2:-2].strip()
                index += 1
            else:
                index += 1
                formula_lines: list[str] = []
                while index < len(lines) and lines[index].strip() != "$$":
                    formula_lines.append(lines[index])
                    index += 1
                if index >= len(lines):
                    raise ValueError("Unclosed display-math block")
                index += 1
                formula = "\n".join(formula_lines).strip()
            blocks.append(Block(kind="math", text=formula))
            continue

        if index + 1 < len(lines) and "|" in line and is_table_separator(lines[index + 1]):
            header = split_table_row(line)
            aligns = table_alignments(lines[index + 1])
            index += 2
            rows = [header]
            while index < len(lines) and lines[index].strip() and "|" in lines[index]:
                rows.append(split_table_row(lines[index]))
                index += 1
            width = len(header)
            if any(len(row) != width for row in rows):
                raise ValueError("Markdown table has inconsistent column counts")
            blocks.append(Block(kind="table", rows=rows, alignments=aligns))
            continue

        list_match = LIST_RE.match(line)
        if list_match:
            indent = len(list_match.group(1).replace("\t", "    "))
            ordered = list_match.group(2) is not None
            content = [list_match.group(4)]
            index += 1
            while index < len(lines):
                candidate = lines[index]
                if not candidate.strip():
                    break
                next_match = LIST_RE.match(candidate)
                if next_match or starts_special(lines, index):
                    break
                leading = len(candidate) - len(candidate.lstrip(" "))
                if leading <= indent:
                    break
                content.append(candidate)
                index += 1
            blocks.append(Block(kind="list", text=smart_join(content), level=indent // 3, ordered=ordered))
            continue

        image_match = IMAGE_RE.match(stripped)
        if image_match:
            blocks.append(Block(kind="image", text=image_match.group(1), path=image_match.group(2)))
            index += 1
            continue

        paragraph_lines = [line]
        index += 1
        while index < len(lines) and not starts_special(lines, index):
            paragraph_lines.append(lines[index])
            index += 1
        blocks.append(Block(kind="paragraph", text=smart_join(paragraph_lines)))

    if not title:
        raise ValueError("Markdown must contain a level-1 paper title")
    return title, blocks


def read_utf8_text(path: Path, role: str) -> str:
    raw = path.read_bytes()
    if b"\x00" in raw:
        raise ValueError(f"{role} is not a UTF-8 text file: {path}")
    if raw.startswith(b"\xef\xbb\xbf"):
        raw = raw[3:]
    try:
        return raw.decode("utf-8")
    except UnicodeDecodeError as exc:
        raise ValueError(f"{role} must use UTF-8 encoding: {path}") from exc


def reject_private_home_paths(text: str, role: str) -> None:
    match = PRIVATE_HOME_RE.search(text)
    if match:
        raise ValueError(f"{role} contains a user-specific absolute path: {match.group(0)}")


def resolve_project_file(raw_path: str, base: Path, project_root: Path, role: str) -> tuple[Path, str]:
    if not raw_path.strip():
        raise ValueError(f"{role} path cannot be empty")
    candidate = Path(raw_path)
    if not candidate.is_absolute():
        candidate = base / candidate
    candidate = candidate.resolve()
    root = project_root.resolve()
    try:
        relative = candidate.relative_to(root)
    except ValueError as exc:
        raise ValueError(f"{role} must stay inside the project root: {candidate}") from exc
    if not candidate.is_file():
        raise FileNotFoundError(candidate)
    return candidate, relative.as_posix()


def parse_markdown_fragment(path: Path, project_root: Path) -> ContentSlotResult:
    text = read_utf8_text(path, "Markdown content fragment")
    reject_private_home_paths(text, "Markdown content fragment")
    if not text.strip():
        raise ValueError(f"Markdown content fragment cannot be empty: {path}")
    _, blocks = parse_markdown("# __CONTENT_FRAGMENT__\n\n" + text)
    for block in blocks:
        if block.kind == "placeholder":
            raise ValueError(f"Nested placeholder {block.text} is forbidden in content fragment {path}")
        if block.kind == "image":
            image, _ = resolve_project_file(block.path, path.parent, project_root, "Fragment image")
            block.path = str(image)
    return ContentSlotResult(
        blocks=blocks,
        inputs=[
            {
                "kind": "markdown",
                "path": path.resolve().relative_to(project_root.resolve()).as_posix(),
                "sha256": sha256_file(path),
                "bytes": path.stat().st_size,
            }
        ],
    )


def validate_ai_header_lines(value: Any) -> list[str]:
    if not isinstance(value, list) or len(value) != 2:
        raise ValueError("Submission source manifest must define exactly two AI header lines")
    lines = [str(item).strip() for item in value]
    if any(not line.startswith("# ") for line in lines):
        raise ValueError("Each AI header line must be a source-code comment beginning with '# '")
    combined = " ".join(lines)
    for required in ("OpenAI Codex", "GPT-5", "OpenAI", "版本发布日期"):
        if required not in combined:
            raise ValueError(f"AI header is missing required disclosure: {required}")
    if not re.search(r"版本发布日期[：:]\s*\d{4}-\d{2}-\d{2}", combined):
        raise ValueError("AI header must contain a verified YYYY-MM-DD model release date")
    return lines


def materialize_content_slot(
    token: str,
    mapping: dict[str, Any],
    base: Path,
    project_root: Path,
) -> ContentSlotResult:
    if not isinstance(mapping, dict):
        raise ValueError(f"Content mapping {token} must be an object")
    mapping_type = str(mapping.get("type") or "").strip()
    if mapping_type == "markdown":
        unknown = sorted(set(mapping) - {"type", "path"})
        if unknown:
            raise ValueError(f"Unsupported keys in Markdown mapping {token}: {', '.join(unknown)}")
        path, _ = resolve_project_file(str(mapping.get("path") or ""), base, project_root, "Markdown content fragment")
        if path.suffix.lower() != ".md":
            raise ValueError(f"Markdown content fragment must use the .md extension: {path}")
        return parse_markdown_fragment(path, project_root)

    if mapping_type != "source_code_appendix":
        raise ValueError(f"Unsupported content mapping type for {token}: {mapping_type or '<empty>'}")
    if token != "SOURCE_CODE_APPENDIX":
        raise ValueError("source_code_appendix type is reserved for SOURCE_CODE_APPENDIX")
    unknown = sorted(set(mapping) - {"type", "manifest"})
    if unknown:
        raise ValueError(f"Unsupported keys in source appendix mapping: {', '.join(unknown)}")
    manifest_path, manifest_relative = resolve_project_file(
        str(mapping.get("manifest") or ""), base, project_root, "Submission source manifest"
    )
    manifest = json.loads(read_utf8_text(manifest_path, "Submission source manifest"))
    if not isinstance(manifest, dict) or manifest.get("schema_version") != "1.0":
        raise ValueError("Submission source manifest must use schema_version 1.0")
    if manifest.get("status") != "frozen":
        raise ValueError("Submission source manifest status must be frozen")
    submission_root_value = str(manifest.get("submission_root") or "")
    submission_root = (project_root / submission_root_value).resolve()
    try:
        submission_root.relative_to(project_root.resolve())
    except ValueError as exc:
        raise ValueError("Submission source root must stay inside the project root") from exc
    if submission_root.name != "提交源码" or manifest_path.parent != submission_root:
        raise ValueError("Submission source manifest must be stored directly in the 提交源码 directory")
    header_lines = validate_ai_header_lines(manifest.get("ai_header_lines"))
    files = manifest.get("files")
    if not isinstance(files, list) or not files:
        raise ValueError("Submission source manifest requires a non-empty files list")

    result = ContentSlotResult(
        inputs=[
            {
                "kind": "source_manifest",
                "path": manifest_relative,
                "sha256": sha256_file(manifest_path),
                "bytes": manifest_path.stat().st_size,
            }
        ]
    )
    seen_sources: set[Path] = set()
    seen_submissions: set[Path] = set()
    for index, item in enumerate(files, start=1):
        if not isinstance(item, dict):
            raise ValueError(f"Submission source manifest item {index} must be an object")
        source_path, source_relative = resolve_project_file(
            str(item.get("source_path") or ""), project_root, project_root, f"Original source item {index}"
        )
        submission_path, submission_relative = resolve_project_file(
            str(item.get("submission_path") or ""), project_root, project_root, f"Submission source item {index}"
        )
        try:
            submission_path.relative_to(submission_root)
        except ValueError as exc:
            raise ValueError(f"Submission source must stay in 提交源码: {submission_relative}") from exc
        if source_path in seen_sources or submission_path in seen_submissions:
            raise ValueError(f"Duplicate source appendix path at item {index}")
        seen_sources.add(source_path)
        seen_submissions.add(submission_path)
        if submission_path.stat().st_size > MAX_SOURCE_FILE_BYTES:
            raise ValueError(f"Submission source exceeds {MAX_SOURCE_FILE_BYTES} bytes: {submission_relative}")
        if sha256_file(source_path) != str(item.get("source_sha256") or "").upper():
            raise ValueError(f"Original source hash changed after submission copy freeze: {source_relative}")
        if sha256_file(submission_path) != str(item.get("submission_sha256") or "").upper():
            raise ValueError(f"Submission source hash mismatch: {submission_relative}")
        if source_path.stat().st_size != int(item.get("source_bytes", -1)):
            raise ValueError(f"Original source byte count mismatch: {source_relative}")
        if submission_path.stat().st_size != int(item.get("submission_bytes", -1)):
            raise ValueError(f"Submission source byte count mismatch: {submission_relative}")
        default_language = SOURCE_LANGUAGE_BY_SUFFIX.get(submission_path.suffix.lower())
        language = str(item.get("language") or default_language or "").strip()
        if not default_language:
            raise ValueError(f"Unsupported submission source extension: {submission_relative}")
        if not language or not re.fullmatch(r"[A-Za-z0-9_+.-]+", language):
            raise ValueError(f"Invalid source language for {submission_relative}: {language!r}")
        code = read_utf8_text(submission_path, "Submission source file")
        reject_private_home_paths(code, "Submission source file")
        if not code.strip():
            raise ValueError(f"Submission source file cannot be empty: {submission_relative}")
        code = code.replace("\r\n", "\n").replace("\r", "\n")
        if code.splitlines()[:2] != header_lines:
            raise ValueError(f"Submission source is missing the exact two-line AI header: {submission_relative}")
        title = str(item.get("title") or f"程序 {index}").strip()
        if not title:
            raise ValueError(f"Submission source title cannot be empty: {submission_relative}")
        result.blocks.extend(
            [
                Block(kind="heading", text=title, level=4),
                Block(kind="paragraph", text=f"源文件：{source_relative}"),
                Block(kind="code", text=code, language=language),
            ]
        )
        result.sources.append(
            {
                "title": title,
                "source_path": source_relative,
                "source_sha256": sha256_file(source_path),
                "path": submission_relative,
                "language": language,
                "sha256": sha256_file(submission_path),
                "bytes": submission_path.stat().st_size,
            }
        )
    return result


def expand_content_slots(
    blocks: Sequence[Block],
    content_map: dict[str, Any],
    content_base: Path,
    project_root: Path,
    state: BuildState,
    *,
    allow_unresolved_errors: bool = False,
) -> list[Block]:
    expanded: list[Block] = []
    seen_slots: set[str] = set()
    for block in blocks:
        if block.kind != "placeholder" or block.text not in content_map:
            expanded.append(block)
            continue
        token = block.text
        if token in seen_slots:
            raise ValueError(f"Content slot appears more than once in Markdown: {token}")
        seen_slots.add(token)
        try:
            materialized = materialize_content_slot(token, content_map[token], content_base, project_root)
        except Exception as exc:
            if not allow_unresolved_errors:
                raise
            state.content_errors.append({"slot": token, "error": str(exc)})
            expanded.append(block)
            continue
        expanded.extend(materialized.blocks)
        state.resolved_content_slots.append(token)
        state.content_inputs.extend({"slot": token, **item} for item in materialized.inputs)
        state.included_sources.extend({"slot": token, **item} for item in materialized.sources)
    return expanded


def collect_formulas(blocks: Iterable[Block], abstract: str) -> list[str]:
    formulas: list[str] = []

    def add_inline(value: str) -> None:
        formulas.extend(match.group(1).strip() for match in INLINE_MATH_RE.finditer(value))

    add_inline(abstract)
    for block in blocks:
        if block.kind == "math":
            formulas.append(block.text.strip())
        elif block.kind in {"paragraph", "heading", "list"}:
            add_inline(block.text)
        elif block.kind == "table":
            for row in block.rows:
                for cell in row:
                    add_inline(cell)
    return formulas


class PandocMathCache:
    def __init__(self, pandoc: Path, work_dir: Path):
        self.pandoc = pandoc
        self.work_dir = work_dir
        self._math: dict[str, etree._Element] = {}

    @staticmethod
    def key(formula: str) -> str:
        return re.sub(r"\s+", " ", formula.strip())

    @staticmethod
    def pandoc_formula(formula: str) -> str:
        # Word math suppresses leading/trailing spaces inside \text{...}.
        # Move those spaces outside as explicit math spacing commands.
        return re.sub(r"\\text\{\s+([^{}]*?)\s+\}", r"\\;\\text{\1}\\;", formula)

    def build(self, formulas: Iterable[str]) -> None:
        unique: list[str] = []
        seen: set[str] = set()
        for formula in formulas:
            key = self.key(formula)
            if key and key not in seen:
                seen.add(key)
                unique.append(key)
        if not unique:
            return

        with tempfile.TemporaryDirectory(prefix="docx-math-", dir=str(self.work_dir)) as tmp_name:
            tmp = Path(tmp_name)
            source = tmp / "math.md"
            output = tmp / "math.docx"
            records = [f"MATH{idx:06d} ${self.pandoc_formula(formula)}$" for idx, formula in enumerate(unique)]
            source.write_text("\n\n".join(records) + "\n", encoding="utf-8")
            command = [
                str(self.pandoc),
                "--from=markdown+tex_math_dollars",
                "--to=docx",
                "--output",
                str(output),
                str(source),
            ]
            completed = subprocess.run(command, capture_output=True, text=True, encoding="utf-8", errors="replace")
            if completed.returncode != 0 or not output.exists():
                raise RuntimeError(f"Pandoc OMML conversion failed: {completed.stderr.strip()}")

            converted = Document(str(output))
            found: dict[int, etree._Element] = {}
            for paragraph in converted.paragraphs:
                marker = re.search(r"MATH(\d{6})", paragraph.text)
                if not marker:
                    continue
                nodes = paragraph._p.findall(f".//{{{M_NS}}}oMath")
                if nodes:
                    found[int(marker.group(1))] = copy.deepcopy(nodes[0])
            missing = sorted(set(range(len(unique))) - set(found))
            if missing:
                details = completed.stderr.strip()
                raise RuntimeError(f"Pandoc did not emit OMML for formula indexes {missing}; {details}")
            self._math = {formula: found[idx] for idx, formula in enumerate(unique)}

    def get(self, formula: str) -> etree._Element:
        key = self.key(formula)
        if key not in self._math:
            raise KeyError(f"Formula not converted to OMML: {formula}")
        return copy.deepcopy(self._math[key])


def set_east_asia_font(style_or_run: Any, east_asia: str, latin: str) -> None:
    style_or_run.font.name = latin
    element = style_or_run._element
    rpr = element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_style(doc: Any, name: str, *, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 12,
              bold: bool = False, italic: bool = False, align: WD_ALIGN_PARAGRAPH | None = None,
              first_indent: float = 0, left_indent: float = 0, hanging: float = 0,
              before: float = 0, after: float = 0, keep_next: bool = False) -> Any:
    styles = doc.styles
    style = styles[name] if name in styles else styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = styles["Normal"]
    set_east_asia_font(style, east_asia, latin)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    fmt = style.paragraph_format
    fmt.alignment = align
    fmt.line_spacing = 1.0
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.first_line_indent = Pt(first_indent) if first_indent else None
    fmt.left_indent = Pt(left_indent) if left_indent else None
    if hanging:
        fmt.first_line_indent = Pt(-hanging)
    fmt.keep_with_next = keep_next
    fmt.widow_control = True
    return style


def configure_styles(doc: Any) -> None:
    normal = doc.styles["Normal"]
    set_east_asia_font(normal, "宋体", "Times New Roman")
    normal.font.size = Pt(12)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 1.0

    add_style(doc, "PaperSpacer", size=2, after=0)
    add_style(doc, "PaperTitle", east_asia="黑体", latin="Arial", size=15, align=WD_ALIGN_PARAGRAPH.CENTER, before=5, after=5, keep_next=True)
    add_style(doc, "PaperAbstractHeading", east_asia="黑体", latin="Arial", size=14, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2, keep_next=True)
    add_style(doc, "PaperAbstract", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=24, after=0)
    add_style(doc, "PaperKeywords", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0)
    add_style(doc, "PaperBody", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=24, after=0)
    add_style(doc, "PaperBodyNoIndent", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0)
    add_style(doc, "PaperHeading1", east_asia="黑体", latin="Arial", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=8, after=4, keep_next=True)
    add_style(doc, "PaperHeading2", east_asia="黑体", latin="Arial", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=2, keep_next=True)
    add_style(doc, "PaperHeading3", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=4, after=1, keep_next=True)
    add_style(doc, "PaperBackmatterHeading", size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=8, after=3, keep_next=True)
    add_style(doc, "PaperList", size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, after=0)
    add_style(doc, "PaperFormula", east_asia="宋体", latin="Cambria Math", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=2)
    add_style(doc, "PaperTableText", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    add_style(doc, "PaperFigure", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, after=0, keep_next=True)
    add_style(doc, "PaperCaption", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, before=2, after=5)
    add_style(doc, "PaperReference", size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, left_indent=21, hanging=21, after=0)
    code_style = add_style(
        doc,
        "PaperCode",
        east_asia="等线",
        latin="Consolas",
        size=7,
        align=WD_ALIGN_PARAGRAPH.LEFT,
        after=0,
    )
    code_style.paragraph_format.line_spacing = Pt(8)
    code_style.paragraph_format.widow_control = False
    code_ppr = code_style.element.get_or_add_pPr()
    snap_to_grid = code_ppr.find(qn("w:snapToGrid"))
    if snap_to_grid is None:
        snap_to_grid = OxmlElement("w:snapToGrid")
        code_ppr.append(snap_to_grid)
    snap_to_grid.set(qn("w:val"), "0")
    add_style(doc, "PaperQaMarker", size=10.5, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=2, after=2)


def set_run_font(run: Any, east_asia: str = "宋体", latin: str = "Times New Roman", size: float | None = None) -> None:
    set_east_asia_font(run, east_asia, latin)
    if size is not None:
        run.font.size = Pt(size)


def append_inline(paragraph: Any, text: str, math_cache: PandocMathCache, *, force_bold: bool = False) -> None:
    token_re = re.compile(r"(\*\*.+?\*\*|`[^`]+`|(?<!\\)\$(?!\$).+?(?<!\\)\$)", re.DOTALL)
    position = 0
    for match in token_re.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position : match.start()])
            run.bold = force_bold
        token = match.group(0)
        if token.startswith("$"):
            paragraph._p.append(math_cache.get(token[1:-1]))
        elif token.startswith("**"):
            append_inline(paragraph, token[2:-2], math_cache, force_bold=True)
        else:
            run = paragraph.add_run(token[1:-1])
            run.bold = force_bold
            set_run_font(run, "等线", "Consolas", 10.5)
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        run.bold = force_bold


def add_display_math(doc: Any, formula: str, math_cache: PandocMathCache) -> None:
    paragraph = doc.add_paragraph(style="PaperFormula")
    math_para = OxmlElement("m:oMathPara")
    math_para_pr = OxmlElement("m:oMathParaPr")
    justification = OxmlElement("m:jc")
    justification.set(qn("m:val"), "center")
    math_para_pr.append(justification)
    math_para.append(math_para_pr)
    math_para.append(math_cache.get(formula))
    paragraph._p.append(math_para)


def clear_document_body_after_first_table(doc: Any) -> None:
    body = doc._element.body
    first_table = next((child for child in body if child.tag == qn("w:tbl")), None)
    if first_table is None or body.index(first_table) != 0:
        raise ValueError("Official template must begin with the identity table")
    section_properties = body.sectPr
    for child in list(body):
        if child is first_table or child is section_properties:
            continue
        body.remove(child)


def replace_cell_text(cell: Any, value: str, *, latin: bool = False, size: float = 10.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "PaperTableText"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(value)
    set_run_font(run, "宋体", "Times New Roman", size)
    if latin:
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(doc: Any) -> None:
    section = doc.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    for child in list(paragraph._p):
        paragraph._p.remove(child)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    set_run_font(run, "宋体", "Times New Roman", 10.5)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    display = OxmlElement("w:t")
    display.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, display, end):
        run._r.append(element)
    sect_pr = section._sectPr
    page_type = sect_pr.find(qn("w:pgNumType"))
    if page_type is None:
        page_type = OxmlElement("w:pgNumType")
        sect_pr.append(page_type)
    page_type.set(qn("w:start"), "1")


def set_numbering(paragraph: Any, *, ordered: bool, level: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), str(min(max(level, 0), 8)))
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1" if ordered else "2")
    num_pr.append(ilvl)
    num_pr.append(num_id)


def visual_length(text: str) -> int:
    plain = INLINE_MATH_RE.sub("MATH", text)
    plain = re.sub(r"[*_`]", "", plain)
    return max(1, sum(2 if ord(char) > 127 else 1 for char in plain))


def calculate_column_widths(rows: Sequence[Sequence[str]], total: int = USABLE_WIDTH_DXA) -> list[int]:
    columns = len(rows[0])
    weights = []
    for col in range(columns):
        maximum = max(visual_length(row[col]) for row in rows)
        weights.append(min(max(maximum, 6), 28))
    minimum = 720 if columns >= 7 else 900
    if minimum * columns > total:
        minimum = total // columns
    remaining = total - minimum * columns
    weight_sum = sum(weights)
    widths = [minimum + int(remaining * weight / weight_sum) for weight in weights]
    widths[-1] += total - sum(widths)
    return widths


def set_table_borders(table: Any, *, visible: bool) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single" if visible else "nil")
        node.set(qn("w:sz"), "4")
        node.set(qn("w:color"), "000000")
        node.set(qn("w:space"), "0")


def set_exact_table_geometry(table: Any, widths: Sequence[int], *, visible_borders: bool = True) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    cell_margins = tbl_pr.find(qn("w:tblCellMar"))
    if cell_margins is None:
        cell_margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(cell_margins)
    for side, value in (("top", 60), ("left", 100), ("bottom", 60), ("right", 100)):
        node = cell_margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            cell_margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.insert(0, tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Twips(widths[index])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_table_borders(table, visible=visible_borders)


def add_markdown_table(doc: Any, block: Block, math_cache: PandocMathCache) -> None:
    table = doc.add_table(rows=len(block.rows), cols=len(block.rows[0]))
    widths = calculate_column_widths(block.rows)
    set_exact_table_geometry(table, widths)
    for row_index, source_row in enumerate(block.rows):
        row = table.rows[row_index]
        if row_index == 0:
            tr_pr = row._tr.get_or_add_trPr()
            repeat = OxmlElement("w:tblHeader")
            repeat.set(qn("w:val"), "true")
            tr_pr.append(repeat)
        for col_index, value in enumerate(source_row):
            cell = row.cells[col_index]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            paragraph.style = "PaperTableText"
            alignment = block.alignments[col_index] if col_index < len(block.alignments) else "left"
            paragraph.alignment = {
                "left": WD_ALIGN_PARAGRAPH.LEFT,
                "center": WD_ALIGN_PARAGRAPH.CENTER,
                "right": WD_ALIGN_PARAGRAPH.RIGHT,
            }[alignment]
            append_inline(paragraph, value, math_cache, force_bold=row_index == 0)
            if row_index == 0:
                shading = cell._tc.get_or_add_tcPr().find(qn("w:shd"))
                if shading is None:
                    shading = OxmlElement("w:shd")
                    cell._tc.get_or_add_tcPr().append(shading)
                shading.set(qn("w:val"), "clear")
                shading.set(qn("w:fill"), "F2F2F2")


def resolve_figure_path(raw_path: str, base: Path) -> Path:
    candidate = Path(raw_path)
    if not candidate.is_absolute():
        candidate = base / candidate
    return candidate.resolve()


def picture_size_cm(path: Path, max_width: float, max_height: float = 17.0) -> tuple[float, float]:
    with Image.open(path) as image:
        width_px, height_px = image.size
    ratio = height_px / max(width_px, 1)
    width = max_width
    height = width * ratio
    if height > max_height:
        height = max_height
        width = height / max(ratio, 1e-9)
    return width, height


def set_picture_alt(inline_shape: Any, alt: str) -> None:
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt or "论文图像")


def add_single_figure(doc: Any, path: Path, caption: str, alt: str, width_cm: float | None = None) -> None:
    width, _ = picture_size_cm(path, min(width_cm or 15.5, 15.8))
    paragraph = doc.add_paragraph(style="PaperFigure")
    shape = paragraph.add_run().add_picture(str(path), width=Cm(width))
    set_picture_alt(shape, alt or caption)
    if caption:
        caption_paragraph = doc.add_paragraph(style="PaperCaption")
        caption_paragraph.add_run(caption)


def add_figure_grid(
    doc: Any,
    items: Sequence[dict[str, Any]],
    base: Path,
    caption: str,
    columns: int | None = None,
) -> list[Path]:
    count = len(items)
    if count == 0:
        raise ValueError("Figure grid requires at least one item")
    if columns is None:
        columns = count
    if isinstance(columns, bool) or not isinstance(columns, int) or not 1 <= columns <= count:
        raise ValueError(f"Figure grid columns must be an integer from 1 to {count}")
    row_count = (count + columns - 1) // columns
    table = doc.add_table(rows=row_count, cols=columns)
    widths = [USABLE_WIDTH_DXA // columns] * columns
    widths[-1] += USABLE_WIDTH_DXA - sum(widths)
    set_exact_table_geometry(table, widths, visible_borders=False)
    for row in table.rows:
        tr_pr = row._tr.get_or_add_trPr()
        cant_split = tr_pr.find(qn("w:cantSplit"))
        if cant_split is None:
            tr_pr.append(OxmlElement("w:cantSplit"))

    remainder = count % columns
    if remainder == 1 and columns > 1:
        last_row = row_count - 1
        table.cell(last_row, 0).merge(table.cell(last_row, columns - 1))

    resolved: list[Path] = []
    for index, item in enumerate(items):
        path = resolve_figure_path(str(item["path"]), base)
        if not path.exists():
            raise FileNotFoundError(path)
        resolved.append(path)
        cell = table.cell(index // columns, index % columns)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        paragraph.style = "PaperFigure"
        max_width = 15.8 / columns - 0.18
        width, _ = picture_size_cm(path, max_width, 8.0)
        shape = paragraph.add_run().add_picture(str(path), width=Cm(width))
        set_picture_alt(shape, str(item.get("alt") or item.get("label") or caption))
        label = str(item.get("label") or "").strip()
        if label:
            label_paragraph = cell.add_paragraph(style="PaperCaption")
            label_paragraph.add_run(label)

    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_with_next = True
    if caption:
        caption_paragraph = doc.add_paragraph(style="PaperCaption")
        caption_paragraph.paragraph_format.keep_together = True
        caption_paragraph.add_run(caption)
    return resolved


def add_figure_from_mapping(doc: Any, token: str, mapping: dict[str, Any], base: Path, state: BuildState) -> None:
    if "items" in mapping:
        columns_value = mapping.get("columns")
        if columns_value is not None and (isinstance(columns_value, bool) or not isinstance(columns_value, int)):
            raise ValueError(f"Figure mapping {token} columns must be an integer")
        paths = add_figure_grid(
            doc,
            list(mapping["items"]),
            base,
            str(mapping.get("caption") or ""),
            columns_value,
        )
        state.inserted_figures.extend(str(path) for path in paths)
        return
    path_value = mapping.get("path")
    if not path_value:
        raise ValueError(f"Figure mapping {token} must contain path or items")
    path = resolve_figure_path(str(path_value), base)
    if not path.exists():
        raise FileNotFoundError(path)
    add_single_figure(
        doc,
        path,
        str(mapping.get("caption") or ""),
        str(mapping.get("alt") or mapping.get("caption") or token),
        float(mapping["width_cm"]) if mapping.get("width_cm") is not None else None,
    )
    state.inserted_figures.append(str(path))


def heading_style(text: str, markdown_level: int) -> str:
    if text in {"AI 工具使用声明", "参考文献", "附录"}:
        return "PaperBackmatterHeading"
    if markdown_level <= 2:
        return "PaperHeading1"
    if markdown_level == 3:
        return "PaperHeading2"
    return "PaperHeading3"


def add_code_block(doc: Any, code: str) -> None:
    normalized = code.replace("\r\n", "\n").replace("\r", "\n").rstrip("\n")
    paragraph = doc.add_paragraph(style="PaperCode")
    paragraph.paragraph_format.widow_control = False
    ppr = paragraph._p.get_or_add_pPr()
    snap_to_grid = ppr.find(qn("w:snapToGrid"))
    if snap_to_grid is None:
        snap_to_grid = OxmlElement("w:snapToGrid")
        ppr.append(snap_to_grid)
    snap_to_grid.set(qn("w:val"), "0")
    run = paragraph.add_run(normalized or " ")
    set_run_font(run, "等线", "Consolas", 7)
    shading = ppr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        ppr.append(shading)
    shading.set(qn("w:fill"), "F7F7F7")


def add_blocks(
    doc: Any,
    blocks: Sequence[Block],
    math_cache: PandocMathCache,
    figure_map: dict[str, Any],
    figure_base: Path,
    markdown_base: Path,
    internal_qa: bool,
    state: BuildState,
) -> None:
    first_body_heading = True
    in_references = False
    for block in blocks:
        if block.kind == "heading":
            paragraph = doc.add_paragraph(style=heading_style(block.text, block.level))
            if first_body_heading:
                paragraph.paragraph_format.page_break_before = True
                first_body_heading = False
            if block.text == "附录":
                paragraph.paragraph_format.page_break_before = True
                in_references = False
            elif block.text == "参考文献":
                in_references = True
            append_inline(paragraph, block.text, math_cache)
        elif block.kind == "paragraph":
            paragraph = doc.add_paragraph(style="PaperReference" if in_references else "PaperBody")
            append_inline(paragraph, block.text, math_cache)
        elif block.kind == "math":
            add_display_math(doc, block.text, math_cache)
        elif block.kind == "list":
            paragraph = doc.add_paragraph(style="PaperList")
            set_numbering(paragraph, ordered=block.ordered, level=block.level)
            append_inline(paragraph, block.text, math_cache)
            if block.ordered:
                state.numbered_paragraphs += 1
            else:
                state.bullet_paragraphs += 1
        elif block.kind == "table":
            add_markdown_table(doc, block, math_cache)
        elif block.kind == "code":
            add_code_block(doc, block.text)
        elif block.kind == "image":
            path = resolve_figure_path(block.path, markdown_base)
            if not path.exists():
                raise FileNotFoundError(path)
            add_single_figure(doc, path, block.text, block.text)
            state.inserted_figures.append(str(path))
        elif block.kind == "placeholder":
            token = block.text
            if token in figure_map:
                add_figure_from_mapping(doc, token, figure_map[token], figure_base, state)
            else:
                state.unresolved.append(token)
                if internal_qa:
                    paragraph = doc.add_paragraph(style="PaperQaMarker")
                    paragraph.add_run(f"[INTERNAL-QA: unresolved source slot {token}]")
        else:
            raise ValueError(f"Unsupported Markdown block kind: {block.kind}")


def numbering_xml() -> bytes:
    root = etree.Element(f"{{{W_NS}}}numbering", nsmap={"w": W_NS})
    for abstract_id, ordered in ((0, True), (1, False)):
        abstract = etree.SubElement(root, f"{{{W_NS}}}abstractNum")
        abstract.set(f"{{{W_NS}}}abstractNumId", str(abstract_id))
        etree.SubElement(abstract, f"{{{W_NS}}}multiLevelType").set(f"{{{W_NS}}}val", "multilevel")
        for level in range(9):
            lvl = etree.SubElement(abstract, f"{{{W_NS}}}lvl")
            lvl.set(f"{{{W_NS}}}ilvl", str(level))
            etree.SubElement(lvl, f"{{{W_NS}}}start").set(f"{{{W_NS}}}val", "1")
            etree.SubElement(lvl, f"{{{W_NS}}}numFmt").set(f"{{{W_NS}}}val", "decimal" if ordered else "bullet")
            if ordered:
                label = ".".join(f"%{idx + 1}" for idx in range(level + 1)) + "."
            else:
                label = "•" if level % 3 == 0 else ("○" if level % 3 == 1 else "▪")
            etree.SubElement(lvl, f"{{{W_NS}}}lvlText").set(f"{{{W_NS}}}val", label)
            etree.SubElement(lvl, f"{{{W_NS}}}lvlJc").set(f"{{{W_NS}}}val", "left")
            ppr = etree.SubElement(lvl, f"{{{W_NS}}}pPr")
            tabs = etree.SubElement(ppr, f"{{{W_NS}}}tabs")
            tab = etree.SubElement(tabs, f"{{{W_NS}}}tab")
            tab.set(f"{{{W_NS}}}val", "num")
            tab.set(f"{{{W_NS}}}pos", str(720 + level * 480))
            indent = etree.SubElement(ppr, f"{{{W_NS}}}ind")
            indent.set(f"{{{W_NS}}}left", str(720 + level * 480))
            indent.set(f"{{{W_NS}}}hanging", "360")
            if not ordered:
                rpr = etree.SubElement(lvl, f"{{{W_NS}}}rPr")
                fonts = etree.SubElement(rpr, f"{{{W_NS}}}rFonts")
                fonts.set(f"{{{W_NS}}}ascii", "宋体")
                fonts.set(f"{{{W_NS}}}hAnsi", "宋体")
                fonts.set(f"{{{W_NS}}}eastAsia", "宋体")
    for num_id, abstract_id in ((1, 0), (2, 1)):
        num = etree.SubElement(root, f"{{{W_NS}}}num")
        num.set(f"{{{W_NS}}}numId", str(num_id))
        etree.SubElement(num, f"{{{W_NS}}}abstractNumId").set(f"{{{W_NS}}}val", str(abstract_id))
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def xml_bytes(root: etree._Element) -> bytes:
    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def next_relationship_id(root: etree._Element) -> str:
    ids = []
    for relationship in root.findall(f"{{{REL_NS}}}Relationship"):
        match = re.fullmatch(r"rId(\d+)", relationship.get("Id", ""))
        if match:
            ids.append(int(match.group(1)))
    return f"rId{max(ids, default=0) + 1}"


def remove_custom_properties(parts: dict[str, bytes]) -> None:
    parts.pop("docProps/custom.xml", None)
    if "_rels/.rels" in parts:
        root = etree.fromstring(parts["_rels/.rels"])
        for rel in list(root.findall(f"{{{REL_NS}}}Relationship")):
            if (rel.get("Target") or "").endswith("docProps/custom.xml"):
                root.remove(rel)
        parts["_rels/.rels"] = xml_bytes(root)
    if "[Content_Types].xml" in parts:
        root = etree.fromstring(parts["[Content_Types].xml"])
        for override in list(root.findall(f"{{{CT_NS}}}Override")):
            if override.get("PartName") == "/docProps/custom.xml":
                root.remove(override)
        parts["[Content_Types].xml"] = xml_bytes(root)


def scrub_metadata(parts: dict[str, bytes]) -> None:
    if "docProps/core.xml" in parts:
        root = etree.fromstring(parts["docProps/core.xml"])
        clear_names = {"creator", "lastModifiedBy", "title", "subject", "description", "keywords", "category"}
        for element in root.iter():
            if etree.QName(element).localname in clear_names:
                element.text = ""
        revision = root.find(f"{{{CP_NS}}}revision")
        if revision is not None:
            revision.text = "1"
        parts["docProps/core.xml"] = xml_bytes(root)
    if "docProps/app.xml" in parts:
        root = etree.fromstring(parts["docProps/app.xml"])
        for name in ("Company", "Manager", "Template", "HyperlinkBase"):
            element = root.find(f"{{{EP_NS}}}{name}")
            if element is not None:
                element.text = ""
        parts["docProps/app.xml"] = xml_bytes(root)
    remove_custom_properties(parts)


def clean_settings(parts: dict[str, bytes]) -> None:
    path = "word/settings.xml"
    root = etree.fromstring(parts[path])
    for tag in ("rsids", "docVars"):
        for element in root.findall(f"{{{W_NS}}}{tag}"):
            root.remove(element)
    update = root.find(f"{{{W_NS}}}updateFields")
    if update is None:
        update = etree.SubElement(root, f"{{{W_NS}}}updateFields")
    update.set(f"{{{W_NS}}}val", "true")
    parts[path] = xml_bytes(root)


def strip_rsid_attributes(parts: dict[str, bytes]) -> None:
    story = re.compile(r"word/(?:document|header\d+|footer\d+|footnotes|endnotes)\.xml$")
    for path in list(parts):
        if not story.fullmatch(path):
            continue
        root = etree.fromstring(parts[path])
        changed = False
        for element in root.iter():
            for attribute in list(element.attrib):
                name = etree.QName(attribute)
                if name.namespace == W_NS and name.localname.startswith("rsid"):
                    del element.attrib[attribute]
                    changed = True
        if changed:
            parts[path] = xml_bytes(root)


def add_numbering_part(parts: dict[str, bytes]) -> None:
    parts["word/numbering.xml"] = numbering_xml()
    content_types = etree.fromstring(parts["[Content_Types].xml"])
    if not any(node.get("PartName") == "/word/numbering.xml" for node in content_types.findall(f"{{{CT_NS}}}Override")):
        override = etree.SubElement(content_types, f"{{{CT_NS}}}Override")
        override.set("PartName", "/word/numbering.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml")
    parts["[Content_Types].xml"] = xml_bytes(content_types)

    rel_path = "word/_rels/document.xml.rels"
    relationships = etree.fromstring(parts[rel_path])
    rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering"
    if not any(node.get("Type") == rel_type for node in relationships.findall(f"{{{REL_NS}}}Relationship")):
        relationship = etree.SubElement(relationships, f"{{{REL_NS}}}Relationship")
        relationship.set("Id", next_relationship_id(relationships))
        relationship.set("Type", rel_type)
        relationship.set("Target", "numbering.xml")
    parts[rel_path] = xml_bytes(relationships)


def postprocess_package(raw_docx: Path, output_docx: Path, *, needs_numbering: bool) -> None:
    with zipfile.ZipFile(raw_docx, "r") as archive:
        parts = {info.filename: archive.read(info.filename) for info in archive.infolist()}
    if needs_numbering:
        add_numbering_part(parts)
    clean_settings(parts)
    scrub_metadata(parts)
    strip_rsid_attributes(parts)

    temp_output = output_docx.with_name(output_docx.name + ".tmp")
    with zipfile.ZipFile(temp_output, "w", compression=zipfile.ZIP_DEFLATED, compresslevel=9) as archive:
        for name in sorted(parts):
            archive.writestr(name, parts[name])
    os.replace(temp_output, output_docx)


def load_figure_map(path: Path | None) -> tuple[dict[str, Any], Path]:
    if path is None:
        return {}, Path.cwd()
    data = json.loads(path.read_text(encoding="utf-8"))
    if not isinstance(data, dict):
        raise ValueError("Figure map must be a JSON object keyed by placeholder token")
    mappings = {key: value for key, value in data.items() if not key.startswith("_")}
    invalid = sorted(key for key in mappings if not TOKEN_RE.fullmatch(key))
    if invalid:
        raise ValueError("Invalid figure-map placeholder tokens: " + ", ".join(invalid))
    return mappings, path.parent.resolve()


def load_content_map(path: Path | None) -> tuple[dict[str, Any], Path]:
    if path is None:
        return {}, Path.cwd()
    data = json.loads(read_utf8_text(path, "Content map"))
    if not isinstance(data, dict):
        raise ValueError("Content map must be a JSON object keyed by placeholder token")
    mappings = {key: value for key, value in data.items() if not key.startswith("_")}
    invalid = sorted(key for key in mappings if not TOKEN_RE.fullmatch(key))
    if invalid:
        raise ValueError("Invalid content-map placeholder tokens: " + ", ".join(invalid))
    for token, mapping in mappings.items():
        if not isinstance(mapping, dict):
            raise ValueError(f"Content mapping {token} must be an object")
    return mappings, path.parent.resolve()


def load_metadata(path: Path | None) -> dict[str, Any]:
    if path is None:
        return {}
    data = yaml.safe_load(path.read_text(encoding="utf-8")) or {}
    if not isinstance(data, dict):
        raise ValueError("Paper metadata file must be a YAML/JSON mapping")
    unknown = sorted(set(data) - {"abstract", "keywords"})
    if unknown:
        raise ValueError("Unsupported paper metadata keys: " + ", ".join(unknown))
    return data


def parse_keywords(value: Any) -> list[str]:
    if value is None:
        return []
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    return [item.strip() for item in re.split(r"[；;,，]", str(value)) if item.strip()]


def validate_identity(group: str, competition_id: str, internal_qa: bool, output: Path, project_root: Path) -> None:
    if internal_qa:
        if group != QA_IDENTITY or competition_id != QA_IDENTITY:
            raise ValueError("Internal QA requires --group INTERNAL-QA and --competition-id INTERNAL-QA")
        tmp_root = (project_root / "tmp").resolve()
        try:
            output.resolve().relative_to(tmp_root)
        except ValueError as exc:
            raise ValueError("Internal QA output must be inside the project tmp directory") from exc
        lower_name = output.name.lower()
        if "qa" not in lower_name and "internal" not in lower_name:
            raise ValueError("Internal QA output filename must visibly contain qa or internal")
        if re.fullmatch(r"acm[0-9]{7}\.(?:docx|pdf)", lower_name):
            raise ValueError("Internal QA output cannot use a formal submission-style filename")
        return
    if group not in ALLOWED_GROUPS:
        raise ValueError(f"Formal group must be one of {sorted(ALLOWED_GROUPS)}")
    if not CM_PATTERN.fullmatch(competition_id):
        raise ValueError("Formal competition ID must match CM followed by exactly seven digits")


def build(args: argparse.Namespace) -> dict[str, Any]:
    script_path = Path(__file__).resolve()
    project_root = script_path.parents[2]
    template = Path(args.template).resolve()
    markdown = Path(args.markdown).resolve()
    output = Path(args.output).resolve()
    validate_identity(args.group, args.competition_id, args.internal_qa, output, project_root)
    if output == template:
        raise ValueError("Output path must differ from the retained official template")
    if output.suffix.lower() != ".docx":
        raise ValueError("Output must use the .docx extension")
    if output.exists() and not args.force:
        raise FileExistsError(f"Output already exists; pass --force to replace: {output}")
    if sha256_file(template) != TEMPLATE_SHA256:
        raise ValueError("Official template SHA-256 mismatch; refresh the template audit before building")

    raw_markdown = markdown.read_text(encoding="utf-8")
    front_matter, body_markdown = parse_front_matter(raw_markdown)
    title, blocks = parse_markdown(body_markdown)
    metadata_path = Path(args.metadata_file).resolve() if args.metadata_file else None
    metadata = load_metadata(metadata_path)
    abstract = args.abstract or metadata.get("abstract") or front_matter.get("abstract") or ""
    if args.abstract_file:
        abstract = Path(args.abstract_file).read_text(encoding="utf-8").strip()
    keyword_source = args.keywords if args.keywords is not None else metadata.get("keywords", front_matter.get("keywords"))
    keywords = parse_keywords(keyword_source)
    if args.internal_qa:
        abstract = abstract or "INTERNAL-QA：摘要尚未冻结。本文件仅用于验证官方模板、公式、表格、图像、页码和渲染链，不得作为竞赛提交文件。"
        keywords = keywords or ["INTERNAL-QA", "模板保真", "渲染检查"]
    if not abstract.strip():
        raise ValueError("Formal build requires an abstract via front matter, --abstract, or --abstract-file")
    if not keywords:
        raise ValueError("Formal build requires keywords via front matter or --keywords")

    figure_map_path = Path(args.figure_map).resolve() if args.figure_map else None
    figure_map, figure_base = load_figure_map(figure_map_path)
    content_map_path = Path(args.content_map).resolve() if args.content_map else None
    content_map, content_base = load_content_map(content_map_path)
    overlap = sorted(set(figure_map) & set(content_map))
    if overlap:
        raise ValueError("Placeholder tokens cannot appear in both figure and content maps: " + ", ".join(overlap))
    if not args.internal_qa:
        for token, item in figure_map.items():
            encoded = json.dumps(item, ensure_ascii=False).lower()
            if token.startswith("FIGURE_Q4") and "preview" in encoded:
                raise ValueError("Formal Q4 figure mapping cannot point to a preview asset")

    pandoc_path = Path(args.pandoc or shutil.which("pandoc") or "")
    if not pandoc_path.exists():
        raise FileNotFoundError("Pandoc is required for native OMML conversion")

    output.parent.mkdir(parents=True, exist_ok=True)
    state = BuildState()
    blocks = expand_content_slots(
        blocks,
        content_map,
        content_base,
        project_root,
        state,
        allow_unresolved_errors=args.internal_qa,
    )
    formulas = collect_formulas(blocks, abstract)
    math_cache = PandocMathCache(pandoc_path.resolve(), output.parent)
    math_cache.build(formulas)
    state.formula_count = len({math_cache.key(item) for item in formulas if item.strip()})

    doc = Document(str(template))
    configure_styles(doc)
    clear_document_body_after_first_table(doc)
    identity_table = doc.tables[0]
    replace_cell_text(identity_table.cell(1, 0), args.group, size=9 if not args.internal_qa else 8.5)
    replace_cell_text(identity_table.cell(1, 2), args.competition_id, latin=True, size=10.5 if not args.internal_qa else 8.5)

    doc.add_paragraph(style="PaperSpacer")
    title_paragraph = doc.add_paragraph(style="PaperTitle")
    title_paragraph.add_run(title)
    abstract_heading = doc.add_paragraph(style="PaperAbstractHeading")
    abstract_heading.add_run("摘要")
    for paragraph_text in re.split(r"\n\s*\n", abstract.strip()):
        paragraph = doc.add_paragraph(style="PaperAbstract")
        append_inline(paragraph, smart_join(paragraph_text.splitlines()), math_cache)
    keyword_paragraph = doc.add_paragraph(style="PaperKeywords")
    label = keyword_paragraph.add_run("关键词：")
    label.bold = True
    keyword_paragraph.add_run("；".join(keywords))

    add_blocks(
        doc,
        blocks,
        math_cache,
        figure_map,
        figure_base,
        markdown.parent,
        args.internal_qa,
        state,
    )
    if state.unresolved and not args.internal_qa:
        raise ValueError("Unresolved content/figure slots: " + ", ".join(state.unresolved))

    add_page_number(doc)
    doc.core_properties.author = ""
    doc.core_properties.last_modified_by = ""
    doc.core_properties.title = ""
    doc.core_properties.subject = ""
    doc.core_properties.keywords = ""
    doc.core_properties.comments = ""

    with tempfile.NamedTemporaryFile(prefix="paper-build-", suffix=".docx", dir=output.parent, delete=False) as handle:
        raw_docx = Path(handle.name)
    try:
        doc.save(str(raw_docx))
        postprocess_package(raw_docx, output, needs_numbering=bool(state.numbered_paragraphs or state.bullet_paragraphs))
    finally:
        raw_docx.unlink(missing_ok=True)

    if sha256_file(template) != TEMPLATE_SHA256:
        raise RuntimeError("Retained official template changed during build")

    completed = subprocess.run([str(pandoc_path), "--version"], capture_output=True, text=True, errors="replace")
    pandoc_version = (completed.stdout.splitlines() or [""])[0]
    manifest = {
        "schema_version": "1.0",
        "mode": "internal_qa" if args.internal_qa else "formal",
        "template": str(template),
        "template_sha256": TEMPLATE_SHA256,
        "markdown": str(markdown),
        "markdown_sha256": sha256_file(markdown),
        "metadata_file": str(metadata_path) if metadata_path else None,
        "metadata_sha256": sha256_file(metadata_path) if metadata_path else None,
        "content_map": str(content_map_path) if content_map_path else None,
        "content_map_sha256": sha256_file(content_map_path) if content_map_path else None,
        "output": str(output),
        "output_sha256": sha256_file(output),
        "group": args.group,
        "competition_id": args.competition_id,
        "pandoc": pandoc_version,
        "omml_formula_count": state.formula_count,
        "numbered_paragraphs": state.numbered_paragraphs,
        "bullet_paragraphs": state.bullet_paragraphs,
        "inserted_figures": state.inserted_figures,
        "resolved_content_slots": state.resolved_content_slots,
        "content_inputs": state.content_inputs,
        "included_source_files": state.included_sources,
        "content_errors": state.content_errors,
        "unresolved_slots": state.unresolved,
        "formal_submission_allowed": not args.internal_qa and not state.unresolved,
    }
    manifest_path = Path(args.manifest).resolve() if args.manifest else output.with_suffix(".build.json")
    if manifest_path.exists() and not args.force:
        raise FileExistsError(f"Manifest already exists; pass --force to replace: {manifest_path}")
    manifest_path.parent.mkdir(parents=True, exist_ok=True)
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    return manifest


def default_paths() -> tuple[Path, Path]:
    project_root = Path(__file__).resolve().parents[2]
    return (
        project_root / "00_赛题与附件" / "2026第七届华数杯竞赛论文模板.docx",
        project_root / "论文" / "内容稿.md",
    )


def parse_args() -> argparse.Namespace:
    template, markdown = default_paths()
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--template", default=str(template), help="Retained official DOCX template")
    parser.add_argument("--markdown", default=str(markdown), help="Paper Markdown source")
    parser.add_argument("--output", required=True, help="Output DOCX path, different from the template")
    parser.add_argument("--group", required=True, help="Formal group or INTERNAL-QA in QA mode")
    parser.add_argument("--competition-id", required=True, help="Formal CMxxxxxxx ID or INTERNAL-QA in QA mode")
    parser.add_argument("--abstract", help="Abstract text; front matter is used when omitted")
    parser.add_argument("--abstract-file", help="UTF-8 abstract file, overriding --abstract/front matter")
    parser.add_argument("--keywords", help="Semicolon/comma-separated keywords; front matter is used when omitted")
    parser.add_argument("--metadata-file", help="YAML/JSON with abstract and keywords; CLI values override it")
    parser.add_argument("--figure-map", help="JSON mapping from HTML placeholder tokens to image assets")
    parser.add_argument("--content-map", help="JSON mapping from HTML placeholder tokens to UTF-8 fragments/source allowlists")
    parser.add_argument("--pandoc", help="Pandoc executable; PATH lookup is used when omitted")
    parser.add_argument("--manifest", help="Build manifest JSON path; defaults beside output")
    parser.add_argument("--internal-qa", action="store_true", help="Allow unresolved slots only for a tmp INTERNAL-QA build")
    parser.add_argument("--force", action="store_true", help="Replace an existing output and manifest")
    return parser.parse_args()


def main() -> int:
    try:
        manifest = build(parse_args())
    except Exception as exc:
        print(f"[build_docx] ERROR: {exc}", flush=True)
        return 2
    print(json.dumps(manifest, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
